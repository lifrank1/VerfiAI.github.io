import sys
import json
import os
import re
import traceback
from collections import defaultdict

# Try importing optional dependencies with fallbacks
try:
    import fitz  # PyMuPDF for PDFs
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file with improved layout preservation."""
    if fitz is None:
        return "Error: PyMuPDF not installed. Run: pip install PyMuPDF"
    
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page_num, page in enumerate(doc):
            # Extract text with better layout preservation
            text += f"\n--- Page {page_num+1} ---\n"
            text += page.get_text("text") + "\n"
            
            # Extract annotations if available
            annots = page.annots()
            if annots:
                text += "\n--- Annotations ---\n"
                for annot in annots:
                    text += f"{annot.info['content']}\n" if 'content' in annot.info else ""
    except Exception as e:
        return f"Error extracting PDF text: {e}"
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file with improved structure preservation."""
    if docx is None:
        return "Error: python-docx not installed. Run: pip install python-docx"
    
    text = ""
    try:
        doc = docx.Document(docx_path)
        
        # Extract document properties if available
        prop_names = ['title', 'author', 'subject', 'keywords', 'comments']
        props = {}
        for prop in prop_names:
            if hasattr(doc.core_properties, prop):
                prop_value = getattr(doc.core_properties, prop)
                if prop_value:
                    props[prop] = prop_value
        
        if props:
            text += "--- Document Properties ---\n"
            for key, value in props.items():
                text += f"{key.capitalize()}: {value}\n"
            text += "\n"
        
        # Extract headers
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name.startswith('Heading'):
                text += f"\n{paragraph.text}\n"
            else:
                text += paragraph.text + "\n"
                
        # Extract tables
        for table in doc.tables:
            text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text += " | ".join(row_text) + "\n"
            text += "\n"
    except Exception as e:
        return f"Error extracting DOCX text: {e}"
    return text

def extract_text_from_txt(txt_path):
    """Extract text from a TXT file."""
    try:
        with open(txt_path, "r", encoding="utf-8") as file:
            return file.read()
    except UnicodeDecodeError:
        # Try different encodings if UTF-8 fails
        encodings = ['latin-1', 'iso-8859-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(txt_path, "r", encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        return "Error: Could not decode file with any supported encoding"
    except Exception as e:
        return f"Error reading TXT file: {e}"

def extract_references(text):
    """
    Enhanced reference extraction with support for multiple citation styles.
    Detects:
    - Numbered references (e.g., '[1]', '1.', etc.)
    - Author-year citations (e.g., 'Smith et al., 2020')
    - Bibliography entries with common patterns
    """
    references = []
    
    # Pattern 1: Numbered references with brackets [1], [2], etc.
    bracket_pattern = re.compile(r"\n\s*\[(\d{1,3})\]\s*(.+?)(?=\n\s*\[\d{1,3}\]|$)", re.DOTALL)
    bracket_matches = bracket_pattern.findall(text)
    for match in bracket_matches:
        ref_number, ref_content = match
        ref_content = ref_content.strip().replace("\n", " ")
        references.append(f"[{ref_number}] {ref_content}")
    
    # Pattern 2: Numbered references with period 1., 2., etc.
    period_pattern = re.compile(r"\n\s*(\d{1,3})\.\s*(.+?)(?=\n\s*\d{1,3}\.|$)", re.DOTALL)
    period_matches = period_pattern.findall(text)
    for match in period_matches:
        ref_number, ref_content = match
        ref_content = ref_content.strip().replace("\n", " ")
        references.append(f"{ref_number}. {ref_content}")
    
    # Pattern 3: Author-year style references
    # Look for common patterns in bibliography sections
    bib_sections = re.findall(r"(?:References|Bibliography|Works Cited|Literature Cited).*?(?=\n\s*(?:[A-Z][a-z]+|\d|\[|$))", text, re.DOTALL)
    
    if bib_sections:
        bib_text = max(bib_sections, key=len)  # Use the longest matching section
        
        # Split by common entry patterns (author names typically start entries)
        entries = re.split(r"\n\s*(?=[A-Z][a-zA-Z\-]+,|\([0-9]{4}\))", bib_text)
        
        for entry in entries:
            entry = entry.strip()
            if len(entry) > 20:  # Minimum length to be considered a reference
                references.append(entry.replace("\n", " "))
    
    # Remove duplicates while preserving order
    unique_references = []
    seen = set()
    for ref in references:
        normalized = re.sub(r'\s+', ' ', ref).strip()
        if normalized not in seen and len(normalized) > 20:
            seen.add(normalized)
            unique_references.append(ref)
    
    return unique_references

def extract_metadata(text):
    """Extract metadata like title, authors, abstract, and keywords from the document."""
    metadata = {}
    
    # Try to extract title (usually at the beginning, in larger font or centered)
    title_patterns = [
        r"(?:^|\n\n)([A-Z][A-Za-z0-9\s:,\-–—]+)(?:\n\n|\n[A-Z])",  # All caps or title case at start
        r"(?:TITLE|Title):\s*([^\n]+)",  # Explicit title marker
    ]
    
    for pattern in title_patterns:
        title_match = re.search(pattern, text[:1000])  # Look only at the beginning
        if title_match:
            metadata["title"] = title_match.group(1).strip()
            break
    
    # Try to extract authors
    author_patterns = [
        r"(?:^|\n)(?:Authors?|BY):\s*([^\n]+)",  # Explicit author marker
        r"(?<=\n\n)([A-Z][a-z]+(?:\s[A-Z][a-z]+)*(?:,\s[A-Z][a-z]+(?:\s[A-Z][a-z]+)*)+)(?=\n\n)",  # Name patterns
    ]
    
    for pattern in author_patterns:
        author_match = re.search(pattern, text[:2000])
        if author_match:
            metadata["authors"] = [a.strip() for a in author_match.group(1).split(",") if a.strip()]
            break
    
    # Try to extract abstract
    abstract_match = re.search(r"(?:Abstract|ABSTRACT):\s*(.*?)(?:\n\n|\n[A-Z]|Keywords:|$)", text, re.DOTALL)
    if abstract_match:
        metadata["abstract"] = abstract_match.group(1).strip().replace("\n", " ")
    
    # Try to extract keywords
    keywords_match = re.search(r"(?:Keywords|KEYWORDS):\s*(.*?)(?:\n\n|$)", text, re.DOTALL)
    if keywords_match:
        keywords = keywords_match.group(1).strip()
        metadata["keywords"] = [k.strip() for k in re.split(r"[,;]", keywords) if k.strip()]
    
    return metadata

def analyze_citation_patterns(text):
    """Analyze in-text citation patterns to identify citation style."""
    # Count occurrences of different citation patterns
    patterns = {
        "numbered_brackets": r"\[\d+\]",                    # [1], [2,3], etc.
        "numbered_parentheses": r"\(\d+\)",                 # (1), (2,3), etc.
        "author_year": r"\([A-Za-z]+(?:\set\sal\.)?(?:,\s\d{4}|\s\d{4})\)", # (Smith, 2020), (Smith et al., 2020)
        "superscript": r"(?<=[a-zA-Z])\d+(?:,\d+)*(?=[,\.\s])",  # superscript numbers
    }
    
    counts = defaultdict(int)
    for style, pattern in patterns.items():
        counts[style] = len(re.findall(pattern, text))
    
    # Determine the most likely citation style
    if counts:
        most_common = max(counts.items(), key=lambda x: x[1])
        if most_common[1] > 0:
            return {
                "style": most_common[0],
                "count": most_common[1],
                "all_counts": dict(counts)
            }
    
    return {"style": "unknown", "count": 0, "all_counts": dict(counts)}

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file path provided"}))
        return

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(json.dumps({"error": "File not found"}))
        return

    file_ext = os.path.splitext(file_path)[1].lower()
    extracted_text = ""

    try:
        if file_ext == ".pdf":
            extracted_text = extract_text_from_pdf(file_path)
        elif file_ext == ".docx":
            extracted_text = extract_text_from_docx(file_path)
        elif file_ext == ".txt":
            extracted_text = extract_text_from_txt(file_path)
        else:
            print(json.dumps({"error": "Unsupported file format"}))
            return

        # Extract various components
        references = extract_references(extracted_text)
        metadata = extract_metadata(extracted_text)
        citation_analysis = analyze_citation_patterns(extracted_text)
        
        # Prepare output
        output = {
            "text": extracted_text,
            "references": references,
            "metadata": metadata,
            "citation_style": citation_analysis["style"],
            "file_type": file_ext[1:],  # Remove the dot
            "file_name": os.path.basename(file_path)
        }
        
        print(json.dumps(output))
    except Exception as e:
        error_details = traceback.format_exc()
        print(json.dumps({
            "error": f"Error processing file: {str(e)}",
            "details": error_details
        }))

if __name__ == "__main__":
    main()
